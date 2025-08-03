import os
import datetime
from docx import Document
from app.utils.watermark import get_watermark_text

def export_to_docx(title: str, content: str, filename: str, subheading: str = None, folder="app/memory"):
    doc = Document()
    doc.add_heading(title, level=1)

    if subheading:
        doc.add_heading(subheading, level=2)

    doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("\n" + content)
    doc.add_paragraph("\n---\n" + get_watermark_text())

    # Ensure memory folder exists
    os.makedirs(folder, exist_ok=True)

    full_path = os.path.join(folder, filename)
    doc.save(full_path)
    return full_path